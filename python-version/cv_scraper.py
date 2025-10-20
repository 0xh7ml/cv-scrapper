#!/usr/bin/env python3
"""
Enhanced CV Scraper - Extract data from PDF and DOCX CV files using Google Gemini AI
Features: Rate limiting, retry logic, batch processing, progress tracking, data validation
"""

import argparse
import asyncio
import csv
import json
import logging
import os
import re
import shutil
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass, asdict, field
from pathlib import Path
from typing import List, Optional, Dict, Any, Tuple
import zipfile
import xml.etree.ElementTree as ET
import time
import threading
from datetime import datetime
from enum import Enum

# Import with graceful error handling for optional dependencies
try:
    import google.generativeai as genai
except ImportError:
    genai = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    import pypdf
except ImportError:
    pypdf = None

try:
    from tqdm import tqdm
except ImportError:
    tqdm = None


# Configure logging with rotating file handler
def setup_logging(verbose: bool = False, log_file: str = 'cv_scraper.log'):
    """Setup logging configuration with file rotation"""
    log_level = logging.DEBUG if verbose else logging.INFO
    
    # Create formatters
    file_formatter = logging.Formatter(
        '%(asctime)s - %(name)s - %(levelname)s - [%(filename)s:%(lineno)d] - %(message)s'
    )
    console_formatter = logging.Formatter(
        '%(asctime)s - %(levelname)s - %(message)s',
        datefmt='%H:%M:%S'
    )
    
    # File handler
    file_handler = logging.FileHandler(log_file, encoding='utf-8')
    file_handler.setLevel(logging.DEBUG)
    file_handler.setFormatter(file_formatter)
    
    # Console handler
    console_handler = logging.StreamHandler()
    console_handler.setLevel(log_level)
    console_handler.setFormatter(console_formatter)
    
    # Configure root logger
    root_logger = logging.getLogger()
    root_logger.setLevel(logging.DEBUG)
    root_logger.handlers = []  # Clear existing handlers
    root_logger.addHandler(file_handler)
    root_logger.addHandler(console_handler)
    
    return logging.getLogger(__name__)


logger = logging.getLogger(__name__)


class ProcessingStatus(Enum):
    """Status of CV processing"""
    PENDING = "pending"
    PROCESSING = "processing"
    SUCCESS = "success"
    FAILED = "failed"
    RETRYING = "retrying"


# Thread-safe counter for CV IDs
class ThreadSafeCounter:
    """Thread-safe counter for generating unique CV IDs"""
    def __init__(self, start=1):
        self._value = start
        self._lock = threading.Lock()
    
    def get_next(self):
        with self._lock:
            current = self._value
            self._value += 1
            return current
    
    def get_current(self):
        with self._lock:
            return self._value - 1


class RateLimiter:
    """Rate limiter for API calls"""
    def __init__(self, max_calls: int, time_window: float):
        """
        Initialize rate limiter
        
        Args:
            max_calls: Maximum number of calls allowed
            time_window: Time window in seconds
        """
        self.max_calls = max_calls
        self.time_window = time_window
        self.calls = []
        self.lock = asyncio.Lock()
    
    async def acquire(self):
        """Acquire permission to make an API call"""
        async with self.lock:
            now = time.time()
            # Remove old calls outside the time window
            self.calls = [call_time for call_time in self.calls if now - call_time < self.time_window]
            
            if len(self.calls) >= self.max_calls:
                sleep_time = self.time_window - (now - self.calls[0])
                if sleep_time > 0:
                    logger.debug(f"Rate limit reached, sleeping for {sleep_time:.2f}s")
                    await asyncio.sleep(sleep_time)
                    return await self.acquire()
            
            self.calls.append(now)


@dataclass
class CVData:
    """Data class to hold extracted CV information"""
    cv_id: str
    name: str = ""
    phone: str = ""
    email: str = ""
    education: str = ""
    skills: str = ""
    experience: str = ""
    linkedin: str = ""
    location: str = ""
    summary: str = ""
    certifications: str = ""
    languages: str = ""
    confidence_score: float = 0.0
    extraction_time: float = 0.0
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary for CSV writing"""
        return asdict(self)
    
    def is_valid(self) -> bool:
        """Check if CV has at least some useful data"""
        critical_fields = [self.name, self.phone, self.email]
        return any(bool(field.strip()) for field in critical_fields)
    
    def get_completeness_score(self) -> float:
        """Calculate completeness score (0-1) based on filled fields"""
        fields = [self.name, self.phone, self.email, self.education, 
                 self.skills, self.experience, self.linkedin, self.location,
                 self.summary, self.certifications, self.languages]
        filled = sum(1 for field in fields if field and field.strip())
        return filled / len(fields)


@dataclass
class ProcessingResult:
    """Result of processing a single CV file"""
    cv_data: Optional[CVData]
    original_filename: str
    new_filename: str
    success: bool
    status: ProcessingStatus = ProcessingStatus.SUCCESS
    error: Optional[str] = None
    retry_count: int = 0
    processing_time: float = 0.0


class DataValidator:
    """Validates and normalizes extracted CV data"""
    
    EMAIL_REGEX = re.compile(r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$')
    PHONE_REGEX = re.compile(r'[\d\s\-\+\(\)]{8,}')
    
    @staticmethod
    def validate_email(email: str) -> str:
        """Validate and normalize email address"""
        if not email:
            return ""
        
        email = email.strip().lower()
        if DataValidator.EMAIL_REGEX.match(email):
            return email
        
        # Try to extract email if it's embedded in text
        emails = re.findall(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', email)
        return emails[0].lower() if emails else ""
    
    @staticmethod
    def validate_phone(phone: str) -> str:
        """Validate and normalize phone number"""
        if not phone:
            return ""
        
        # Extract digits and common phone characters
        phone = re.sub(r'[^\d\s\-\+\(\)]', '', phone.strip())
        
        # Check if it looks like a phone number
        if DataValidator.PHONE_REGEX.search(phone):
            return phone
        
        return ""
    
    @staticmethod
    def normalize_name(name: str) -> str:
        """Normalize person's name"""
        if not name:
            return ""
        
        # Remove extra whitespace and title case
        name = ' '.join(name.strip().split())
        
        # Remove common prefixes
        prefixes = ['mr.', 'mrs.', 'ms.', 'dr.', 'prof.']
        name_lower = name.lower()
        for prefix in prefixes:
            if name_lower.startswith(prefix):
                name = name[len(prefix):].strip()
                break
        
        return name.title() if name else ""
    
    @staticmethod
    def truncate_field(text: str, max_length: int = 500) -> str:
        """Truncate field to maximum length"""
        if not text:
            return ""
        
        text = text.strip()
        if len(text) <= max_length:
            return text
        
        return text[:max_length-3] + "..."
    
    @classmethod
    def validate_cv_data(cls, cv_data: CVData) -> CVData:
        """Validate and normalize all fields in CV data"""
        cv_data.name = cls.normalize_name(cv_data.name)
        cv_data.email = cls.validate_email(cv_data.email)
        cv_data.phone = cls.validate_phone(cv_data.phone)
        cv_data.education = cls.truncate_field(cv_data.education, 500)
        cv_data.skills = cls.truncate_field(cv_data.skills, 500)
        cv_data.experience = cls.truncate_field(cv_data.experience, 1000)
        cv_data.summary = cls.truncate_field(cv_data.summary, 300)
        cv_data.certifications = cls.truncate_field(cv_data.certifications, 300)
        
        return cv_data


class TextExtractor:
    """Handles text extraction from different file formats"""
    
    @staticmethod
    def extract_from_pdf(file_path: Path) -> str:
        """Extract text from PDF using multiple extraction methods"""
        extractors = [
            ("PyMuPDF", TextExtractor._extract_with_pymupdf),
            ("pdfplumber", TextExtractor._extract_with_pdfplumber),
            ("pypdf", TextExtractor._extract_with_pypdf),
            ("PyPDF2", TextExtractor._extract_with_pypdf2),
        ]
        
        best_text = ""
        best_length = 0
        
        for extractor_name, extractor_func in extractors:
            try:
                text = extractor_func(file_path)
                if text and len(text.strip()) > best_length:
                    best_text = text
                    best_length = len(text.strip())
                    logger.debug(f"Successfully extracted {best_length} chars using {extractor_name} for {file_path.name}")
            except Exception as e:
                logger.debug(f"{extractor_name} failed for {file_path}: {e}")
                continue
        
        if not best_text.strip():
            raise Exception("All PDF extractors failed to extract meaningful text")
        
        return best_text
    
    @staticmethod
    def _extract_with_pymupdf(file_path: Path) -> str:
        """Extract text using PyMuPDF (fitz)"""
        if fitz is None:
            raise ImportError("PyMuPDF (fitz) is not installed")
        
        doc = fitz.open(file_path)
        text_parts = []
        
        for page in doc:
            text = page.get_text()
            if text.strip():
                text_parts.append(text)
        
        doc.close()
        return '\n'.join(text_parts)
    
    @staticmethod
    def _extract_with_pdfplumber(file_path: Path) -> str:
        """Extract text using pdfplumber (good for complex layouts)"""
        if pdfplumber is None:
            raise ImportError("pdfplumber is not installed")
        
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text and text.strip():
                    text_parts.append(text)
        
        return '\n'.join(text_parts)
    
    @staticmethod
    def _extract_with_pypdf(file_path: Path) -> str:
        """Extract text using pypdf (newer version of PyPDF2)"""
        if pypdf is None:
            raise ImportError("pypdf is not installed")
        
        with open(file_path, 'rb') as file:
            pdf_reader = pypdf.PdfReader(file)
            text_parts = []
            
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text.strip():
                    text_parts.append(text)
            
            return '\n'.join(text_parts)
    
    @staticmethod
    def _extract_with_pypdf2(file_path: Path) -> str:
        """Extract text using PyPDF2 (basic extraction)"""
        if PyPDF2 is None:
            raise ImportError("PyPDF2 is not installed")
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text_parts = []
            
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text.strip():
                    text_parts.append(text)
            
            return '\n'.join(text_parts)
    
    @staticmethod
    def extract_from_docx(file_path: Path) -> str:
        """Extract text from DOCX using python-docx (fallback to zip method)"""
        try:
            if Document is not None:
                doc = Document(file_path)
                text_parts = []
                
                # Extract paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text)
                
                # Extract tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text_parts.append(cell.text)
                
                return '\n'.join(text_parts)
            else:
                logger.info(f"python-docx not available, using zip method for {file_path}")
                raise ImportError("python-docx not installed")
            
        except Exception as e:
            logger.warning(f"python-docx failed for {file_path}, trying zip method: {e}")
            
            try:
                return TextExtractor._extract_docx_fallback(file_path)
            except Exception as e2:
                raise Exception(f"Both DOCX extractors failed: python-docx({e}), zip({e2})")
    
    @staticmethod
    def _extract_docx_fallback(file_path: Path) -> str:
        """Fallback DOCX extraction using zip and XML parsing"""
        with zipfile.ZipFile(file_path, 'r') as zip_file:
            try:
                doc_xml = zip_file.read('word/document.xml')
                root = ET.fromstring(doc_xml)
                
                text_parts = []
                for elem in root.iter():
                    if elem.tag.endswith('}t') and elem.text:
                        text_parts.append(elem.text)
                
                return ' '.join(text_parts)
                
            except KeyError:
                raise Exception("word/document.xml not found in DOCX file")


class GeminiCVExtractor:
    """Handles CV data extraction using Google Gemini AI with retry logic"""
    
    def __init__(self, api_key: str, rate_limiter: Optional[RateLimiter] = None, max_retries: int = 3):
        """Initialize Gemini client"""
        if genai is None:
            raise ImportError("google-generativeai is not installed. Install it with: pip install google-generativeai")
        
        genai.configure(api_key=api_key)
        self.model = genai.GenerativeModel('gemini-2.0-flash-exp')
        self.rate_limiter = rate_limiter
        self.max_retries = max_retries
        
        self.generation_config = genai.types.GenerationConfig(
            temperature=0.1,
            max_output_tokens=1000,
            top_p=0.8,
            top_k=40
        )
    
    async def extract_cv_data(self, text: str, cv_id: str) -> CVData:
        """Extract CV data from text using Gemini AI with retry logic"""
        for attempt in range(self.max_retries):
            try:
                if self.rate_limiter:
                    await self.rate_limiter.acquire()
                
                prompt = self._create_extraction_prompt(text)
                
                response = await asyncio.get_event_loop().run_in_executor(
                    None,
                    lambda: self.model.generate_content(
                        prompt,
                        generation_config=self.generation_config
                    )
                )
                
                if not response.text:
                    raise Exception("Empty response from Gemini")
                
                json_text = self._clean_json_response(response.text)
                cv_dict = json.loads(json_text)
                
                cv_data = CVData(
                    cv_id=cv_id,
                    name=cv_dict.get('name', ''),
                    phone=cv_dict.get('phone', ''),
                    email=cv_dict.get('email', ''),
                    education=cv_dict.get('education', ''),
                    skills=cv_dict.get('skills', ''),
                    experience=cv_dict.get('experience', ''),
                    linkedin=cv_dict.get('linkedin', ''),
                    location=cv_dict.get('location', ''),
                    summary=cv_dict.get('summary', ''),
                    certifications=cv_dict.get('certifications', ''),
                    languages=cv_dict.get('languages', ''),
                    confidence_score=cv_dict.get('confidence_score', 0.7)
                )
                
                # Validate and normalize data
                cv_data = DataValidator.validate_cv_data(cv_data)
                
                return cv_data
                
            except json.JSONDecodeError as e:
                logger.warning(f"Attempt {attempt + 1}/{self.max_retries} - JSON parse error: {e}")
                if attempt < self.max_retries - 1:
                    await asyncio.sleep(2 ** attempt)  # Exponential backoff
                else:
                    raise Exception(f"Failed to parse Gemini response after {self.max_retries} attempts")
            
            except Exception as e:
                logger.warning(f"Attempt {attempt + 1}/{self.max_retries} - API error: {e}")
                if attempt < self.max_retries - 1:
                    await asyncio.sleep(2 ** attempt)
                else:
                    raise Exception(f"Gemini API error after {self.max_retries} attempts: {e}")
    
    @staticmethod
    def _create_extraction_prompt(text: str) -> str:
        """Create the prompt for CV data extraction"""
        return f"""You are an expert CV/Resume data extraction system. Extract the following information from the CV text and return it as a JSON object.

Required fields (return empty string "" if not found):
- name: Full name of the person
- phone: Phone number (with country code if present)
- email: Email address
- education: Educational background, degrees, institutions (key points only)
- skills: Technical skills, competencies, tools, programming languages (comma-separated)
- experience: Work experience summary - job titles, companies, durations
- linkedin: LinkedIn profile URL
- location: Current location or address
- summary: Professional summary or objective (2-3 sentences max)
- certifications: Professional certifications or licenses
- languages: Languages spoken (comma-separated)
- confidence_score: Your confidence in the extraction quality (0.0 to 1.0)

Extraction rules:
1. Return empty string "" for any field not found
2. Keep fields concise - max 500 chars per field (except experience: 1000 chars)
3. For skills: list technologies, tools, languages comma-separated
4. For experience: focus on recent roles, include years if available
5. Extract URLs exactly as written
6. Return ONLY valid JSON, no markdown, no extra text
7. Normalize phone/email formats
8. Set confidence_score based on data quality and completeness

CV Text:
{text[:8000]}"""  # Limit text to avoid token limits
    
    @staticmethod
    def _clean_json_response(response_text: str) -> str:
        """Clean Gemini response to extract valid JSON"""
        text = response_text.strip()
        
        # Remove markdown code fences
        if text.startswith('```json'):
            text = text[7:]
        elif text.startswith('```'):
            text = text[3:]
        
        if text.endswith('```'):
            text = text[:-3]
        
        text = text.strip()
        
        # Try to extract JSON object if embedded in text
        if not text.startswith('{'):
            start = text.find('{')
            if start != -1:
                text = text[start:]
        
        if not text.endswith('}'):
            end = text.rfind('}')
            if end != -1:
                text = text[:end+1]
        
        return text


class ProgressTracker:
    """Track and display progress of CV processing"""
    
    def __init__(self, total: int, use_tqdm: bool = True):
        self.total = total
        self.processed = 0
        self.successful = 0
        self.failed = 0
        self.lock = threading.Lock()
        self.start_time = time.time()
        self.use_tqdm = use_tqdm and tqdm is not None
        
        if self.use_tqdm:
            self.pbar = tqdm(total=total, desc="Processing CVs", unit="file")
    
    def update(self, success: bool):
        """Update progress"""
        with self.lock:
            self.processed += 1
            if success:
                self.successful += 1
            else:
                self.failed += 1
            
            if self.use_tqdm:
                self.pbar.update(1)
                self.pbar.set_postfix({
                    'success': self.successful,
                    'failed': self.failed
                })
    
    def close(self):
        """Close progress tracker"""
        if self.use_tqdm:
            self.pbar.close()
    
    def get_summary(self) -> str:
        """Get processing summary"""
        elapsed = time.time() - self.start_time
        avg_time = elapsed / self.processed if self.processed > 0 else 0
        
        return (f"\nProcessing Summary:\n"
                f"  Total: {self.total}\n"
                f"  Successful: {self.successful}\n"
                f"  Failed: {self.failed}\n"
                f"  Success Rate: {self.successful/self.total*100:.1f}%\n"
                f"  Total Time: {elapsed:.2f}s\n"
                f"  Average Time/CV: {avg_time:.2f}s")


class CVScraper:
    """Main CV scraper application with enhanced features"""
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx'}
    
    def __init__(self, api_key: str, concurrency: int = 20, max_retries: int = 3, 
                 rate_limit: Tuple[int, float] = (60, 60.0)):
        """
        Initialize CV scraper
        
        Args:
            api_key: Gemini API key
            concurrency: Number of concurrent workers
            max_retries: Maximum retry attempts per file
            rate_limit: Tuple of (max_calls, time_window_seconds)
        """
        self.rate_limiter = RateLimiter(rate_limit[0], rate_limit[1])
        self.gemini_extractor = GeminiCVExtractor(api_key, self.rate_limiter, max_retries)
        self.text_extractor = TextExtractor()
        self.concurrency = concurrency
        self.counter = ThreadSafeCounter()
        self.max_retries = max_retries
    
    def find_and_rename_cv_files(self, input_dir: Path) -> List[Tuple[Path, str]]:
        """Find all supported CV files and rename them incrementally"""
        cv_files = []
        renamed_files = []
        
        # Collect all valid CV files
        for file_path in sorted(input_dir.iterdir()):
            if file_path.is_file() and file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                cv_files.append(file_path)
            elif file_path.is_file():
                logger.debug(f"Skipping unsupported file: {file_path.name}")
        
        logger.info(f"Found {len(cv_files)} CV files to process")
        
        # Rename files incrementally
        for file_path in cv_files:
            cv_id = self.counter.get_next()
            new_filename = f"{cv_id}{file_path.suffix.lower()}"
            new_path = input_dir / new_filename
            
            try:
                if file_path != new_path:  # Only rename if different
                    shutil.move(str(file_path), str(new_path))
                    logger.debug(f"Renamed {file_path.name} -> {new_filename}")
                renamed_files.append((new_path, str(cv_id)))
            except Exception as e:
                logger.error(f"Failed to rename {file_path.name}: {e}")
                continue
        
        return renamed_files
    
    async def process_single_cv_async(self, file_info: Tuple[Path, str], 
                                     progress: Optional[ProgressTracker] = None) -> ProcessingResult:
        """Process a single CV file asynchronously"""
        file_path, cv_id = file_info
        start_time = time.time()
        
        logger.debug(f"Processing CV {cv_id}: {file_path.name}")
        
        try:
            # Extract text based on file extension
            if file_path.suffix.lower() == '.pdf':
                text = await asyncio.get_event_loop().run_in_executor(
                    None, self.text_extractor.extract_from_pdf, file_path
                )
            elif file_path.suffix.lower() == '.docx':
                text = await asyncio.get_event_loop().run_in_executor(
                    None, self.text_extractor.extract_from_docx, file_path
                )
            else:
                raise Exception(f"Unsupported file type: {file_path.suffix}")
            
            if not text.strip():
                raise Exception("No text extracted from file")
            
            logger.debug(f"Extracted {len(text)} characters from CV {cv_id}")
            
            # Extract CV data using Gemini
            cv_data = await self.gemini_extractor.extract_cv_data(text, cv_id)
            cv_data.extraction_time = time.time() - start_time
            
            # Check if we got useful data
            if not cv_data.is_valid():
                if progress:
                    progress.update(False)
                return ProcessingResult(
                    cv_data=None,
                    original_filename=file_path.name,
                    new_filename=file_path.name,
                    success=False,
                    status=ProcessingStatus.FAILED,
                    error="No critical CV data found (name, phone, or email)",
                    processing_time=time.time() - start_time
                )
            
            logger.info(f"Successfully processed CV {cv_id} - Completeness: {cv_data.get_completeness_score():.1%}")
            
            if progress:
                progress.update(True)
            
            return ProcessingResult(
                cv_data=cv_data,
                original_filename=file_path.name,
                new_filename=file_path.name,
                success=True,
                status=ProcessingStatus.SUCCESS,
                processing_time=time.time() - start_time
            )
            
        except Exception as e:
            logger.warning(f"Failed to process CV {cv_id}: {str(e)}")
            if progress:
                progress.update(False)
            
            return ProcessingResult(
                cv_data=None,
                original_filename=file_path.name,
                new_filename=file_path.name,
                success=False,
                status=ProcessingStatus.FAILED,
                error=str(e),
                processing_time=time.time() - start_time
            )
    
    async def process_cvs_concurrently(self, cv_files: List[Tuple[Path, str]], 
                                      show_progress: bool = True) -> Tuple[List[CVData], List[ProcessingResult]]:
        """Process multiple CV files concurrently with progress tracking"""
        logger.info(f"Processing {len(cv_files)} CV files with concurrency: {self.concurrency}")
        
        progress = ProgressTracker(len(cv_files), use_tqdm=show_progress) if show_progress else None
        
        # Create semaphore to limit concurrent operations
        semaphore = asyncio.Semaphore(self.concurrency)
        
        async def process_with_semaphore(file_info):
            async with semaphore:
                return await self.process_single_cv_async(file_info, progress)
        
        # Process all files concurrently
        tasks = [process_with_semaphore(file_info) for file_info in cv_files]
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        if progress:
            progress.close()
        
        # Collect results
        successful_cvs = []
        all_results = []
        
        for result in results:
            if isinstance(result, Exception):
                logger.error(f"Unexpected error: {result}")
                all_results.append(ProcessingResult(
                    cv_data=None,
                    original_filename="unknown",
                    new_filename="unknown",
                    success=False,
                    error=str(result)
                ))
            else:
                all_results.append(result)
                if result.success and result.cv_data:
                    successful_cvs.append(result.cv_data)
        
        if progress:
            logger.info(progress.get_summary())
        else:
            logger.info(f"Processing complete: {len(successful_cvs)}/{len(cv_files)} files processed successfully")
        
        return successful_cvs, all_results
    
    @staticmethod
    def save_to_csv(cv_data: List[CVData], output_file: Path):
        """Save CV data to CSV file with all fields"""
        if not cv_data:
            logger.warning("No CV data to save")
            return
        
        fieldnames = ['cv_id', 'name', 'phone', 'email', 'education', 'skills', 'experience',
                     'linkedin', 'location', 'summary', 'certifications', 'languages',
                     'confidence_score', 'extraction_time']
        
        with open(output_file, 'w', newline='', encoding='utf-8') as csvfile:
            writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
            writer.writeheader()
            
            for cv in cv_data:
                writer.writerow(cv.to_dict())
        
        logger.info(f"Successfully saved {len(cv_data)} CV records to {output_file}")
    
    @staticmethod
    def save_error_report(results: List[ProcessingResult], output_file: Path):
        """Save error report for failed processing attempts"""
        failed_results = [r for r in results if not r.success]
        
        if not failed_results:
            logger.info("No errors to report")
            return
        
        with open(output_file, 'w', newline='', encoding='utf-8') as csvfile:
            fieldnames = ['filename', 'error', 'retry_count', 'processing_time']
            writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
            writer.writeheader()
            
            for result in failed_results:
                writer.writerow({
                    'filename': result.new_filename,
                    'error': result.error or "Unknown error",
                    'retry_count': result.retry_count,
                    'processing_time': f"{result.processing_time:.2f}s"
                })
        
        logger.info(f"Saved error report with {len(failed_results)} failed files to {output_file}")
    
    @staticmethod
    def generate_statistics(cv_data: List[CVData], results: List[ProcessingResult]) -> Dict[str, Any]:
        """Generate processing statistics"""
        if not cv_data:
            return {}
        
        total_time = sum(r.processing_time for r in results)
        successful_count = len(cv_data)
        failed_count = len([r for r in results if not r.success])
        
        completeness_scores = [cv.get_completeness_score() for cv in cv_data]
        confidence_scores = [cv.confidence_score for cv in cv_data]
        
        stats = {
            'total_processed': len(results),
            'successful': successful_count,
            'failed': failed_count,
            'success_rate': successful_count / len(results) if results else 0,
            'total_time': total_time,
            'avg_time_per_cv': total_time / len(results) if results else 0,
            'avg_completeness': sum(completeness_scores) / len(completeness_scores) if completeness_scores else 0,
            'avg_confidence': sum(confidence_scores) / len(confidence_scores) if confidence_scores else 0,
            'min_completeness': min(completeness_scores) if completeness_scores else 0,
            'max_completeness': max(completeness_scores) if completeness_scores else 0,
        }
        
        return stats
    
    @staticmethod
    def save_statistics(stats: Dict[str, Any], output_file: Path):
        """Save processing statistics to JSON file"""
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(stats, f, indent=2)
        
        logger.info(f"Saved statistics to {output_file}")


async def main():
    """Main entry point"""
    parser = argparse.ArgumentParser(
        description="Enhanced CV Scraper - Extract data from PDF and DOCX files using Google Gemini AI",
        formatter_class=argparse.ArgumentDefaultsHelpFormatter
    )
    
    parser.add_argument(
        '-i', '--input',
        type=Path,
        default=Path('./cvs'),
        help='Input directory containing CV files'
    )
    
    parser.add_argument(
        '-o', '--output',
        type=Path,
        default=Path('data.csv'),
        help='Output CSV file'
    )
    
    parser.add_argument(
        '-c', '--concurrency',
        type=int,
        default=20,
        help='Number of concurrent workers'
    )
    
    parser.add_argument(
        '--api-key',
        type=str,
        help='Gemini API key (or set GEMINI_API_KEY env var)'
    )
    
    parser.add_argument(
        '--max-retries',
        type=int,
        default=3,
        help='Maximum retry attempts per file'
    )
    
    parser.add_argument(
        '--rate-limit',
        type=int,
        default=60,
        help='Maximum API calls per minute'
    )
    
    parser.add_argument(
        '--error-report',
        type=Path,
        default=Path('errors.csv'),
        help='Error report output file'
    )
    
    parser.add_argument(
        '--stats',
        type=Path,
        default=Path('stats.json'),
        help='Statistics output file'
    )
    
    parser.add_argument(
        '--no-progress',
        action='store_true',
        help='Disable progress bar'
    )
    
    parser.add_argument(
        '--verbose', '-v',
        action='store_true',
        help='Enable verbose logging'
    )
    
    parser.add_argument(
        '--log-file',
        type=str,
        default='cv_scraper.log',
        help='Log file path'
    )
    
    args = parser.parse_args()
    
    # Setup logging
    global logger
    logger = setup_logging(args.verbose, args.log_file)
    
    logger.info("=" * 60)
    logger.info("Enhanced CV Scraper - Starting")
    logger.info("=" * 60)
    logger.info(f"Input directory: {args.input}")
    logger.info(f"Output file: {args.output}")
    logger.info(f"Concurrency: {args.concurrency}")
    logger.info(f"Max retries: {args.max_retries}")
    logger.info(f"Rate limit: {args.rate_limit} calls/minute")
    
    # Get Gemini API key
    api_key = args.api_key or os.getenv('GEMINI_API_KEY')
    if not api_key:
        logger.error("Gemini API key is required. Use --api-key flag or set GEMINI_API_KEY environment variable")
        return 1
    
    # Create input directory if it doesn't exist
    if not args.input.exists():
        args.input.mkdir(parents=True, exist_ok=True)
        logger.info(f"Created '{args.input}' directory. Please place your CV files there.")
        return 0
    
    # Initialize scraper
    try:
        scraper = CVScraper(
            api_key=api_key,
            concurrency=args.concurrency,
            max_retries=args.max_retries,
            rate_limit=(args.rate_limit, 60.0)
        )
    except ImportError as e:
        logger.error(f"Missing required dependency: {e}")
        logger.error("Install dependencies with: pip install google-generativeai PyMuPDF pdfplumber python-docx")
        return 1
    
    # Find and rename CV files
    start_time = time.time()
    cv_files = scraper.find_and_rename_cv_files(args.input)
    
    if not cv_files:
        logger.info("No valid CV files found in the input directory.")
        return 0
    
    # Process CVs asynchronously
    cv_data, results = await scraper.process_cvs_concurrently(
        cv_files,
        show_progress=not args.no_progress
    )
    
    # Save results
    if cv_data:
        scraper.save_to_csv(cv_data, args.output)
        
        # Save error report
        scraper.save_error_report(results, args.error_report)
        
        # Generate and save statistics
        stats = scraper.generate_statistics(cv_data, results)
        stats['timestamp'] = datetime.now().isoformat()
        stats['input_directory'] = str(args.input)
        stats['output_file'] = str(args.output)
        scraper.save_statistics(stats, args.stats)
        
        # Print final summary
        end_time = time.time()
        processing_time = end_time - start_time
        
        logger.info("=" * 60)
        logger.info("Processing Complete!")
        logger.info("=" * 60)
        logger.info(f"Total CVs processed: {stats['total_processed']}")
        logger.info(f"Successful: {stats['successful']} ({stats['success_rate']*100:.1f}%)")
        logger.info(f"Failed: {stats['failed']}")
        logger.info(f"Total time: {processing_time:.2f} seconds")
        logger.info(f"Average time per CV: {stats['avg_time_per_cv']:.2f} seconds")
        logger.info(f"Average completeness: {stats['avg_completeness']*100:.1f}%")
        logger.info(f"Average confidence: {stats['avg_confidence']*100:.1f}%")
        logger.info("=" * 60)
        logger.info(f"Results saved to: {args.output}")
        logger.info(f"Error report saved to: {args.error_report}")
        logger.info(f"Statistics saved to: {args.stats}")
        logger.info(f"Log file: {args.log_file}")
        
    else:
        logger.warning("No CVs were processed successfully.")
        return 1
    
    return 0


if __name__ == "__main__":
    exit(asyncio.run(main()))